import streamlit as st
import sqlite3
from datetime import datetime
from docx import Document
from io import BytesIO

DB_NAME = "expert2.db"


# --- DB HELPERS ---
def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute(
        '''CREATE TABLE IF NOT EXISTS patients (id INTEGER PRIMARY KEY AUTOINCREMENT, date_naissance TEXT, age INTEGER)''')
    c.execute(
        '''CREATE TABLE IF NOT EXISTS resultats (id INTEGER PRIMARY KEY AUTOINCREMENT, patient_id INTEGER, texte_id INTEGER, reponse TEXT, precision TEXT)''')
    conn.commit()
    conn.close()


def get_categories():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT DISTINCT categorie FROM textes WHERE categorie IS NOT NULL AND categorie != ''")
    cats = [row[0] for row in c.fetchall()]
    conn.close()
    return sorted(cats)


def get_questions(category=None):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    if category:
        c.execute("SELECT id, phrase, categorie, intensite, prec FROM textes WHERE categorie = ?", (category,))
    else:
        c.execute("SELECT id, phrase, categorie, intensite, prec FROM textes")
    questions = c.fetchall()
    conn.close()
    return questions


def get_materiel_eligible(texte_id, age):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    query = """
        SELECT m.nom, m.site FROM materiel m
        JOIN correspondance c ON m.id = c.materiel_id
        WHERE c.texte_id = ? AND m.age_mini <= ? AND m.age_maxi >= ?
    """
    c.execute(query, (texte_id, age, age))
    res = c.fetchall()
    conn.close()
    return res


# --- CALLBACKS ---
def valider_patient(date_naiss):
    today = datetime.now()
    age = today.year - date_naiss.year - ((today.month, today.day) < (date_naiss.month, date_naiss.day))

    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("INSERT INTO patients (date_naissance, age) VALUES (?, ?)", (date_naiss.strftime("%d/%m/%Y"), age))
    st.session_state.patient_id = c.lastrowid
    st.session_state.age = age
    st.session_state.date_naiss = date_naiss.strftime("%d/%m/%Y")
    st.session_state.step = 'selection_mode'  # On passe au choix du mode
    conn.commit()
    conn.close()


def lancer_quiz(category=None):
    st.session_state.questions = get_questions(category)
    st.session_state.selected_cat_name = category if category else "Complet"
    st.session_state.step = 'quiz'


def enregistrer_reponse(q_id, phrase, cat, intensite, rep, precision):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("INSERT INTO resultats (patient_id, texte_id, reponse, precision) VALUES (?, ?, ?, ?)",
              (st.session_state.patient_id, q_id, rep, precision))
    conn.commit()
    conn.close()
    if rep == "Vrai":
        st.session_state.reponses_vrai.append(
            {'id': q_id, 'phrase': phrase, 'categorie': cat, 'intensite': intensite, 'precision': precision})
    st.session_state.current_index += 1


# --- LOGIQUE RAPPORT ---
def generer_rapport_docx(patient_id, date_naiss, age, reponses_vrai):
    doc = Document()
    doc.add_heading(f"Diagnostic Sensoriel", 0)
    doc.add_paragraph(f"Date de naissance : {date_naiss}")
    doc.add_paragraph(f"Âge lors de l'examen : {age} ans")

    doc.add_heading("Récapitulatif des réponses 'Vrai'", level=1)

    if not reponses_vrai:
        doc.add_paragraph("Aucune réponse 'Vrai' enregistrée.")
    else:
        categories = sorted(list(set(r['categorie'] for r in reponses_vrai)))
        materiels_globaux = {}

        for cat in categories:
            doc.add_heading(f"Catégorie : {cat}", level=2)
            items_cat = [r for r in reponses_vrai if r['categorie'] == cat]
            for item in items_cat:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{item['phrase']}").bold = True
                p.add_run(f" (Intensité : {item['intensite']})")
                if item['precision']:
                    doc.add_paragraph(f"   Précision : {item['precision']}")

                mats = get_materiel_eligible(item['id'], age)
                if mats:
                    doc.add_paragraph(f"   Matériel : {', '.join([m[0] for m in mats])}")
                    for m_nom, m_site in mats:
                        materiels_globaux[m_nom] = m_site

        if materiels_globaux:
            doc.add_page_break()
            doc.add_heading("Liste globale du matériel", level=1)
            for nom, site in materiels_globaux.items():
                doc.add_paragraph(f"{nom} (Site : {site})")

    target = BytesIO()
    doc.save(target)
    return target.getvalue()



# --- INTERFACE PRINCIPALE ---
def main():
    st.set_page_config(page_title="Diagnostic Modulaire", page_icon="⚙️")
    init_db()

    if 'step' not in st.session_state: st.session_state.step = 'accueil'
    if 'reponses_vrai' not in st.session_state: st.session_state.reponses_vrai = []
    if 'current_index' not in st.session_state: st.session_state.current_index = 0

    # --- ÉTAPE 1 : ACCUEIL ---
    if st.session_state.step == 'accueil':
        st.title("🩺 Nouveau Diagnostic")
        date_naiss = st.date_input("Date de naissance du patient")
        st.button("Suivant", on_click=valider_patient, args=(date_naiss,))

    # --- ÉTAPE 2 : CHOIX DU MODE (LA NOUVEAUTÉ) ---
    elif st.session_state.step == 'selection_mode':
        st.title("🎯 Type d'analyse")
        st.write(f"Patient de {st.session_state.age} ans. Souhaitez-vous une analyse complète ou ciblée ?")

        col1, col2 = st.columns(2)
        with col1:
            if st.button("🚀 Analyse Complète", use_container_width=True):
                lancer_quiz()
                st.rerun()

        with col2:
            st.write("--- OU ---")
            categories = get_categories()
            choix = st.selectbox("Choisir une catégorie spécifique :", categories)
            if st.button(f"🔍 Analyser : {choix}", use_container_width=True):
                lancer_quiz(choix)
                st.rerun()

    # --- ÉTAPE 3 : QUIZ ---
    elif st.session_state.step == 'quiz':
        idx = st.session_state.current_index
        qs = st.session_state.questions

        if idx < len(qs):
            q_id, phrase, cat, intensite, prec_requise = qs[idx]
            st.progress(idx / len(qs))
            st.subheader(f"Catégorie : {cat}")
            st.markdown(f"#### {phrase}")

            precision = st.text_area("Notes :", key=f"p_{idx}") if prec_requise == "oui" else ""

            c1, c2, c3 = st.columns(3)
            c1.button("✅ Vrai", on_click=enregistrer_reponse, args=(q_id, phrase, cat, intensite, "Vrai", precision))
            c2.button("❌ Faux", on_click=enregistrer_reponse, args=(q_id, phrase, cat, intensite, "Faux", precision))
            c3.button("❓ Ne sais pas", on_click=enregistrer_reponse, args=(q_id, phrase, cat, intensite, "Inconnu", precision))
        else:
            st.session_state.step = 'final'
            st.rerun()

    # --- ÉTAPE 4 : BILAN ---
    elif st.session_state.step == 'final':
        st.title("📊 Résultats")
        st.download_button("📥 Télécharger le Rapport", data=generer_rapport_docx(), file_name="Rapport.docx")
        if st.button("Recommencer"):
            st.session_state.clear()
            st.rerun()


if __name__ == "__main__":
    main()